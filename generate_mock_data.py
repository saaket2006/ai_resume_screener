import os
import docx
from docx.shared import Pt
import random

os.makedirs('data', exist_ok=True)

def create_resume_docx(filename, name, role, email, phone, linkedin, summary, education, skills, projects, certifications, achievements):
    doc = docx.Document()
    
    # Header
    name_run = doc.add_paragraph().add_run(name)
    name_run.font.size = Pt(24)
    name_run.bold = True
    
    role_run = doc.add_paragraph().add_run(role)
    role_run.font.size = Pt(14)
    role_run.bold = True
    
    contact_p = doc.add_paragraph()
    contact_p.add_run(f"📞 {phone} | ✉️ {email} | 🔗 {linkedin}")
    
    doc.add_paragraph() # Spacer
    
    # Sections
    def add_section_header(title):
        p = doc.add_paragraph()
        r = p.add_run(title)
        r.font.size = Pt(14)
        r.bold = True
        # add a small border or just rely on bold
        
    add_section_header("Summary")
    doc.add_paragraph(summary)
    
    add_section_header("Education")
    doc.add_paragraph(education)
    
    add_section_header("Skills")
    for category, skill_list in skills.items():
        p = doc.add_paragraph()
        p.add_run(category + ": ").bold = True
        p.add_run(skill_list)
        
    add_section_header("Projects")
    for proj_title, proj_desc in projects:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(proj_title + " - ").bold = True
        p.add_run(proj_desc)
        
    add_section_header("Certifications")
    for cert in certifications:
        doc.add_paragraph(cert, style='List Bullet')
        
    add_section_header("Achievements & Hackathons")
    for ach in achievements:
        doc.add_paragraph(ach, style='List Bullet')

    doc.save(filename)

# Resume 1: Strong AI/Software Match
create_resume_docx(
    filename='data/resume_1_realistic_ai.docx',
    name='Saaket Baldawa',
    role='AI & Software Engineering Enthusiast',
    email='saaket.hyd@gmail.com',
    phone='+91 8309344858',
    linkedin='linkedin.com/in/saaket-baldawa',
    summary='I am a motivated undergraduate student focused on Generative AI systems and software engineering, with hands-on experience building structured multi-agent pipelines. Strong foundations in Python, APIs, and object-oriented programming with exposure to cloud platforms and modern LLM-based workflows.',
    education='Institute of Aeronautical Engineering\nB.Tech in Computer Science and Engineering (AI & ML) | 2024 - 2028 | CGPA: 9.35',
    skills={
        'Programming': 'Java, Python, React, TypeScript, Object-Oriented Programming (OOP)',
        'AI & ML': 'Artificial Intelligence Fundamentals, Generative AI, LLM-based systems, Multi-Agent Workflows, Machine Learning, scikit-learn, spaCy, Natural Language Processing',
        'Development': 'Python-based application development, application programming interfaces (APIs), FastAPI, Node.js',
        'Tools & Platforms': 'Oracle Cloud Infrastructure, AWS, Docker, Git'
    },
    projects=[
        ('Multi-Agent YouTube Video Summarization System', 'Built an LLM-driven multi-agent system for automated YouTube video summarization, reducing manual viewing time by 80%. Implemented transcript extraction, chunking, and structured summaries via a Streamlit interface.'),
        ('M2A Intelligence Platform', 'Developed an AI platform that converts online meeting recordings into structured notes. Integrated a CUDA-enabled transcription pipeline with a local LLM and modular backend APIs with a React and TypeScript based frontend.')
    ],
    certifications=[
        'Oracle Cloud Infrastructure 2025 Certified AI Foundations Associate',
        'Deloitte Australia - Technology Job Simulation (Virtual Experience)',
        'Python Essentials 2 (PCEP - Certified Associate Python Programmer)'
    ],
    achievements=[
        '5-Day AI Agents Intensive Course with Google (Kaggle)',
        'Winter Coding Contest 5.0 - ACM Student Chapter, VNRVJIET',
        'INCEPTUM - The Beginning of Innovation Hackathon'
    ]
)

# Resume 2: Data Science / Analyst Match
create_resume_docx(
    filename='data/resume_2_realistic_data_science.docx',
    name='Elena Rostova',
    role='Data Scientist',
    email='elena.rostova@example.com',
    phone='+1 555-019-2834',
    linkedin='linkedin.com/in/elenarostova',
    summary='Analytical Data Scientist with 3+ years of experience in statistical analysis, predictive modeling, and data visualization. Adept at transforming complex datasets into actionable business insights using Python and SQL.',
    education='University of Data Science\nM.S. in Data Analytics | 2020 - 2022 | GPA: 3.9',
    skills={
        'Programming': 'Python, R, SQL, MATLAB',
        'Data & Analytics': 'Pandas, NumPy, Tableau, Power BI, Matplotlib, Seaborn',
        'Machine Learning': 'scikit-learn, XGBoost, Random Forests, Linear Regression',
        'Tools': 'Jupyter, Git, Excel'
    },
    projects=[
        ('Customer Churn Prediction Model', 'Developed a random forest classifier to predict customer churn with 85% accuracy, allowing the retention team to target at-risk users.'),
        ('Sales Dashboard Automation', 'Created automated Tableau dashboards connected to an SQL database to visualize real-time regional sales metrics.')
    ],
    certifications=[
        'Google Data Analytics Professional Certificate',
        'AWS Certified Cloud Practitioner'
    ],
    achievements=[
        'Kaggle Competitions Master',
        'Published paper on Predictive Maintenance in Manufacturing'
    ]
)

# Resume 3: Frontend Developer Match
create_resume_docx(
    filename='data/resume_3_realistic_frontend.docx',
    name='Marcus Johnson',
    role='Frontend Web Developer',
    email='mjohnson.frontend@example.com',
    phone='+44 7700 900123',
    linkedin='linkedin.com/in/marcusj-dev',
    summary='Creative Frontend Developer specializing in building responsive, accessible, and highly interactive user interfaces. Passionate about modern JavaScript frameworks and bridging the gap between design and technical implementation.',
    education='Tech Institute of Arts\nB.S. in Web Design and Development | 2018 - 2022',
    skills={
        'Frontend': 'HTML5, CSS3, JavaScript (ES6+), React, Vue.js, Redux',
        'Styling': 'SASS, Tailwind CSS, Bootstrap, Styled Components',
        'Design Tools': 'Figma, Adobe XD, Photoshop',
        'Tools & Workflow': 'Webpack, Vite, Git, npm, Jest'
    },
    projects=[
        ('E-Commerce Component Library', 'Designed and implemented a reusable React component library used across 3 company web platforms to ensure brand consistency.'),
        ('Interactive Portfolio Platform', 'Built a responsive portfolio template utilizing Vue.js and Tailwind CSS with complex scroll animations and dynamic routing.')
    ],
    certifications=[
        'Meta Front-End Developer Professional Certificate',
        'Responsive Web Design Certification (freeCodeCamp)'
    ],
    achievements=[
        'Winner - Best UI/UX Design at Local Hack Day 2023',
        'Open Source Contributor - React Documentation'
    ]
)

print("Realistic DOCX resumes generated successfully.")
